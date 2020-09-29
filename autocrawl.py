import requests
import time
from bs4 import BeautifulSoup
import docx
import re
import sys
import os
from progressbar import progressbar

def getSoup(chap):
    source_code = requests.get('https://truyencv.vn/truyen/legendary-moonlight-sculptor-con-duong-de-vuong/' + str(chap))
    soup = BeautifulSoup(source_code.text, features="lxml")
    return soup

def getTitle(soup):
    title = soup.find('h1', { 'class': "chap-head"})
    return title.string

def getContent(soup):
    content = soup.find('div', { 'class' : "chap-content"})
    return content.p

def getVolNum(title):
    #Kiểm tra Title
    tapRegex = re.compile(r'Tập (\d)(\d)?')
    mo = tapRegex.search(title)
    if mo.group(2) == None:
        volnum = int(mo.group(1))
    else:
        volnum = int(mo.group(1))*10 + int(mo.group(2))
    return volnum

def crawl_vol(vol):

    if vol == 0:
        print("Không có Vol 0")
        exit(1)

    chap_estimate = (vol-1)*10+1
    

    for i in progressbar(range(chap_estimate, chap_estimate + 12), redirect_stdout=True):

            #Lấy source code format qua BeautifulSoup
            soup = getSoup(chap_estimate)

            #Lấy Title crawl về
            title = getTitle(soup)

            #Lấy số Vol về
            volnum = getVolNum(title)

            #Kiểm tra Vol number
            if volnum < vol:
                chap_estimate += 1
                continue
            elif volnum > vol:
                print(" Hoàn Tất")
                break
            else:
                # Mở file doc để bắt đầu ghi
                try:
                    doc = docx.Document('Legendary Moonlight Scupltor Vol ' + str(volnum) + '.docx')
                except:
                    doc = docx.Document()
                    print('Tạo file Legendary Moonlight Scupltor Vol ' + str(volnum))
                #Add title
                doc.add_heading(title)
                print('Bắt đầu viết '+ title)
                
                #Lấy Content về
                content = getContent(soup)
                for line in content:
                    if line.string == None:
                        continue
                    doc.add_paragraph(line.string)
                
                doc.add_page_break()
                #Save lại file
                doc.save('Legendary Moonlight Scupltor Vol ' + str(volnum) + '.docx')
                chap_estimate += 1
os.chdir('C:\\Users\\ThinkKING\\Desktop\\Code\\Automate')
# if len(sys.argv) == 1:
#     print("Please input a number to start crawling")
#     exit(1)
# elif len(sys.argv) > 2:
#     print("Error command line")
#     exit(1)
# else:
#     i = int(sys.argv[1])
for i in range(3,11):
    crawl_vol(i)
